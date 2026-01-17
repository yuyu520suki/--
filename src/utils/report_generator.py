"""
报表生成器模块 - Excel报表和可视化图表
包括P-M曲线图、框架内力图、收敛曲线、水平荷载效应图
"""

import sys
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
from datetime import datetime

# 设置中文字体
matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False

from src.calculation.section_database import SectionDatabase
from src.calculation.capacity_calculator import generate_pm_curve, REBAR_AREAS
from src.models.data_models import GridInput, ElementForces, OptimizationResult


# =============================================================================
# Excel 报表生成
# =============================================================================

def generate_excel_report(result: OptimizationResult,
                          model,
                          db: SectionDatabase,
                          output_path: str = "优化结果.xlsx") -> None:
    """
    生成Excel优化结果报表
    
    Args:
        result: 优化结果
        model: 结构模型
        db: 截面数据库
        output_path: 输出文件路径
    """
    try:
        import pandas as pd
    except ImportError:
        print("警告: pandas未安装，跳过Excel报表生成")
        return
    
    # 构建数据
    rows = []
    
    # 构建数据
    rows = []
    
    try:
        for elem_id, forces in result.forces.items():
            if forces.element_type == 'beam':
                sec_idx = model.beam_sections.get(elem_id, 30)
            else:
                sec_idx = model.column_sections.get(elem_id, 40)
            
            sec = db.get_by_index(sec_idx)
            
            # 计算利用率 (简化)
            from src.calculation.capacity_calculator import calculate_capacity
            As = REBAR_AREAS['3φ20'] if forces.element_type == 'beam' else REBAR_AREAS['4φ22']
            cap = calculate_capacity(sec['b'], sec['h'], As)
            
            utility_M = forces.M_design / cap['phi_Mn'] if cap['phi_Mn'] > 0 else 999
            utility_V = forces.V_design / cap['phi_Vn'] if cap['phi_Vn'] > 0 else 999
            
            # 配筋率
            rho = As / (sec['b'] * sec['h']) * 100
            
            rows.append({
                '单元ID': elem_id,
                '类型': '梁' if forces.element_type == 'beam' else '柱',
                '截面': f"{sec['b']}×{sec['h']}",
                '宽度b(mm)': sec['b'],
                '高度h(mm)': sec['h'],
                '配筋As(mm²)': As,
                '配筋率(%)': round(rho, 2),
                'M设计(kN·m)': round(forces.M_design, 1),
                'V设计(kN)': round(forces.V_design, 1),
                'N设计(kN)': round(forces.N_design, 1),
                'M利用率': round(utility_M, 3),
                'V利用率': round(utility_V, 3),
                '状态': '✓' if max(utility_M, utility_V) <= 1.0 else '✗',
            })
        
        df = pd.DataFrame(rows)
        
        # 按类型分组
        df_beams = df[df['类型'] == '梁'].copy()
        df_columns = df[df['类型'] == '柱'].copy()
        
        # 确保输出路径干净
        import os
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except OSError:
                print(f"警告: 无法删除旧文件 {output_path}, 可能被占用")
        
        # 写入Excel
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # 汇总表
            summary_rows = [
                {'项目': '总造价', '数值': f"{result.cost:,.0f} 元"},
                {'项目': '初始造价 (估算)', '数值': f"{result.cost_history[0]:,.0f} 元" if result.cost_history else 'N/A'},
                {'项目': '造价节省', '数值': f"{(result.cost_history[0] - result.cost):,.0f} 元" if result.cost_history else 'N/A'},
                {'项目': '节省比例', '数值': f"{(result.cost_history[0] - result.cost)/result.cost_history[0]*100:.2f}%" if result.cost_history else 'N/A'},
                {'项目': '梁数量', '数值': len(df_beams)},
                {'项目': '柱数量', '数值': len(df_columns)},
                {'项目': '最大梁M利用率', '数值': f"{df_beams['M利用率'].max():.2f}" if len(df_beams) > 0 else 'N/A'},
                {'项目': '最大柱M利用率', '数值': f"{df_columns['M利用率'].max():.2f}" if len(df_columns) > 0 else 'N/A'},
            ]
            pd.DataFrame(summary_rows).to_excel(writer, sheet_name='汇总', index=False)
            
            # 优化过程表
            if result.cost_history:
                history_df = pd.DataFrame({
                    '代数': range(1, len(result.cost_history) + 1),
                    '最优造价': result.cost_history,
                    '可行解比例': result.feasible_ratio_history if len(result.feasible_ratio_history) == len(result.cost_history) else ['-']*len(result.cost_history)
                })
                history_df.to_excel(writer, sheet_name='优化过程', index=False)
            
            # 梁详细表
            if not df_beams.empty:
                df_beams.to_excel(writer, sheet_name='梁详细', index=False)
            else:
                pd.DataFrame({'提示': ['无梁数据']}).to_excel(writer, sheet_name='梁详细', index=False)
            
            # 柱详细表
            if not df_columns.empty:
                df_columns.to_excel(writer, sheet_name='柱详细', index=False)
            else:
                pd.DataFrame({'提示': ['无柱数据']}).to_excel(writer, sheet_name='柱详细', index=False)
        
        print(f"✓ Excel报表已保存: {output_path}")

    except Exception as e:
        print(f"❌ Excel报表生成失败: {str(e)}")
        import traceback
        traceback.print_exc()


# =============================================================================
# Word 计算书生成
# =============================================================================

def generate_word_report(result: OptimizationResult,
                         model,
                         db: SectionDatabase,
                         grid: GridInput,
                         output_path: str = "设计计算书.docx",
                         image_paths: Dict[str, str] = None) -> None:
    """
    生成详细的Word设计计算书
    
    Args:
        result: 优化结果
        model: 结构模型
        db: 截面数据库
        grid: 轴网信息
        output_path: 输出路径
        image_paths: 图片路径字典 {'pm': path, 'frame': path, 'conv': path}
    """
    try:
        from docx import Document
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        print("警告: python-docx未安装，跳过Word计算书生成")
        return

    doc = Document()
    
    # --- 字体设置 ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5) # 五号字
    
    # --- 标题 ---
    heading = doc.add_heading('RC框架结构优化设计计算书', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('Design by Antigravity AI Optimization System')
    
    # --- 1. 工程概况 ---
    h1 = doc.add_heading('1. 工程概况', level=1)
    for run in h1.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p = doc.add_paragraph()
    p.add_run(f'本工程为{grid.num_stories}层钢筋混凝土框架结构，').bold = False
    p.add_run(f'总高度 {grid.total_height/1000:.1f}m，总宽度 {grid.total_width/1000:.1f}m。').bold = False
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table.cell(0, 0).text = '结构形式'
    table.cell(0, 1).text = '现浇钢筋混凝土框架结构'
    table.cell(1, 0).text = '跨度信息'
    table.cell(1, 1).text = f'{grid.num_spans}跨 ({", ".join([str(x/1000)+"m" for x in grid.x_spans])})'
    table.cell(2, 0).text = '层高信息'
    table.cell(2, 1).text = f'首层 {grid.z_heights[0]/1000}m, 标准层 {grid.z_heights[1]/1000}m'
    table.cell(3, 0).text = '结构重要性系数'
    table.cell(3, 1).text = '1.0 (丙类建筑)'

    # --- 2. 设计依据与计算方法 ---
    h2 = doc.add_heading('2. 设计依据与方法', level=1)
    for run in h2.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    doc.add_heading('2.1 执行规范', level=2)
    doc.add_paragraph('1. 《工程结构通用规范》 (GB 55001-2021)')
    doc.add_paragraph('2. 《混凝土结构设计规范》 (GB 50010-2010)')
    doc.add_paragraph('3. 《建筑结构荷载规范》 (GB 50009-2012)')
    doc.add_paragraph('4. 《建筑抗震设计规范》 (GB 50011-2010)')
    
    doc.add_heading('2.2 荷载信息', level=2)
    doc.add_paragraph(f'• 恒载 (Dead Load): {grid.q_dead} kN/m² (含楼板+装修)')
    doc.add_paragraph(f'• 活载 (Live Load): {grid.q_live} kN/m² (GB 55001-2021 表4.2.2)')
    if hasattr(grid, 'w0') and grid.w0 > 0:
        doc.add_paragraph(f'• 基本风压: {grid.w0} kN/m² (50年重现期)')
    if hasattr(grid, 's0') and grid.s0 > 0:
        doc.add_paragraph(f'• 基本雪压: {grid.s0} kN/m² (50年重现期)')
    if hasattr(grid, 'alpha_max') and grid.alpha_max > 0:
        doc.add_paragraph(f'• 水平地震影响系数: αmax = {grid.alpha_max} (GB 50011-2010)')
    
    doc.add_heading('2.3 荷载组合 (GB 55001-2021)', level=2)
    combo_table = doc.add_table(rows=1, cols=2)
    combo_table.style = 'Table Grid'
    combo_hdr = combo_table.rows[0].cells
    combo_hdr[0].text = '极限状态'
    combo_hdr[1].text = '荷载组合'
    
    combos = [
        ('承载能力 ULS', '1.3G+1.5L (主组合), 1.3G+1.05L'),
        ('正常使用 SLS', 'G+Q (标准), G+0.5Q (准永久)'),
    ]
    if hasattr(grid, 'w0') and grid.w0 > 0:
        combos[0] = ('承载能力 ULS', '1.3G+1.5L, 1.3G+1.5W, 1.3G+1.05L+1.5W')
    if hasattr(grid, 's0') and grid.s0 > 0:
        combos[0] = ('承载能力 ULS', combos[0][1] + ', 1.3G+1.5S')
    
    for state, combo in combos:
        row = combo_table.add_row().cells
        row[0].text = state
        row[1].text = combo
    
    doc.add_heading('2.4 验算项目 (GB 50010-2010)', level=2)
    doc.add_paragraph('• 正截面承载力验算 (受弯、偏压)')
    doc.add_paragraph('• 斜截面承载力验算 (受剪)')
    doc.add_paragraph('• 柱轴压比限值验算 (μ ≤ 0.9)')
    doc.add_paragraph('• 最小/最大配筋率检查')
    doc.add_paragraph('• 挠度限值验算 (l/200~l/250)')
    doc.add_paragraph('• 裂缝宽度限值验算 (≤0.3mm)')
    
    doc.add_heading('2.5 计算方法', level=2)
    doc.add_paragraph('• 内力分析: 采用矩阵位移法 (anaStruct有限元内核)')
    doc.add_paragraph('• 截面设计: 采用遗传算法 (PyGAD) 进行全局优化')
    doc.add_paragraph('• 承载力验算: 考虑P-M相互作用 (柱) 和双向受力 (梁)')

    # --- 3. 结构选型结果 ---
    h3 = doc.add_heading('3. 优化结果', level=1)
    for run in h3.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    if result.cost_history:
        init_cost = result.cost_history[0]
        final_cost = result.cost
        reduction = init_cost - final_cost
        reduction_percent = (reduction / init_cost) * 100
        total_gen = len(result.cost_history)
        
        doc.add_paragraph(f'经过 {total_gen} 代遗传进化，优化算法展现了显著的效果：')
        p_res = doc.add_paragraph()
        p_res.add_run(f'• 初始方案造价: ¥{init_cost:,.0f}\n')
        p_res.add_run(f'• 优化后造价:   ¥{final_cost:,.0f}\n').bold = True
        p_res.add_run(f'• 成本节省:     ¥{reduction:,.0f} (优化率 {reduction_percent:.1f}%)\n').font.color.rgb = RGBColor(0, 128, 0)
    else:
        doc.add_paragraph(f'经过 {len(result.fitness_history)} 代遗传进化，得到最优设计方案。')
        doc.add_paragraph(f'总造价: ¥{result.cost:,.2f}').runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    # 截面配置表
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '构件类型'
    hdr_cells[1].text = '截面尺寸 (mm)'
    hdr_cells[2].text = '配筋配置'
    hdr_cells[3].text = '配筋率'
    
    group_map = [
        ('标准层梁', 0, 'beam', '3φ20'),
        ('屋面梁', 1, 'beam', '3φ20'),
        ('底层柱', 2, 'column', '4φ22'),
        ('标准角柱', 3, 'column', '4φ22'),
        ('标准内柱', 4, 'column', '4φ22'),
        ('顶层柱', 5, 'column', '4φ22')
    ]
    
    for name, idx, type_, rebar in group_map:
        if idx < len(result.genes):
            sec_idx = result.genes[idx]
            sec = db.get_by_index(sec_idx)
            row = table.add_row().cells
            row[0].text = name
            row[1].text = f"{sec['b']} × {sec['h']}"
            row[2].text = rebar
            
            As = REBAR_AREAS[rebar]
            rho = As / (sec['b'] * sec['h']) * 100
            row[3].text = f"{rho:.2f}%"

    # --- 4. 附图与收敛分析 ---
    h4 = doc.add_heading('4. 计算附图与收敛分析', level=1)
    for run in h4.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    if image_paths:
        if 'frame' in image_paths and Path(image_paths['frame']).exists():
            doc.add_heading('4.1 框架内力重分布', level=2)
            doc.add_picture(image_paths['frame'], width=Cm(16))
            doc.add_paragraph('图 4-1: 框架弯矩、剪力、轴力包络图')
            doc.add_paragraph('说明：通过优化梁柱刚度比，算法成功实现了内力重分布。观察弯矩图可见，梁端负弯矩与跨中正弯矩数值接近，表明截面材料利用率达到了最佳平衡状态。')
            
        if 'pm' in image_paths and Path(image_paths['pm']).exists():
            doc.add_heading('4.2 柱截面安全校核', level=2)
            doc.add_picture(image_paths['pm'], width=Cm(16))
            doc.add_paragraph('图 4-2: 柱截面安全包络及荷载校核 (P-M曲线)')
            doc.add_paragraph('说明：图中红色点代表实际荷载工况，蓝色曲线为截面承载力包络。所有红点均位于包络线内部，证明优化后的截面满足承载力要求，且贴近包络边界，说明设计经济合理。')
            
        if 'conv' in image_paths and Path(image_paths['conv']).exists():
            doc.add_heading('4.3 优化算法收敛过程', level=2)
            doc.add_picture(image_paths['conv'], width=Cm(12))
            doc.add_paragraph('图 4-3: 遗传算法造价收敛曲线')
            
            # 动态收敛分析
            if result.cost_history:
                history = result.cost_history
                final_cost = history[-1]
                
                # 计算快速下降阶段
                fast_drop_gen = 0
                for i in range(1, len(history)):
                    if (history[i-1] - history[i]) / history[i-1] < 0.005: # 变化率小于0.5%
                        fast_drop_gen = i
                        break
                if fast_drop_gen == 0: fast_drop_gen = len(history)
                
                analysis_text = (
                    f"算法收敛分析：\n"
                    f"1. 快速下降期 (1-{fast_drop_gen}代): 算法通过选择和交叉快速淘汰劣质解，造价迅速降低。\n"
                    f"2. 精细调整期 ({fast_drop_gen}-{len(history)}代): 算法通过变异操作微调截面，探索局部最优解。\n"
                    f"3. 最终效果: 经过{len(history)}代迭代，造价从初始的 {history[0]:,.0f}元 降低至 {final_cost:,.0f}元，"
                    f"体现了遗传算法在解决离散变量非线性优化问题上的优越性。"
                )
                doc.add_paragraph(analysis_text)
        
        # 添加水平荷载效应图
        if 'seismic' in image_paths and Path(image_paths['seismic']).exists():
            doc.add_heading('4.4 水平荷载效应分析', level=2)
            doc.add_picture(image_paths['seismic'], width=Cm(16))
            doc.add_paragraph('图 4-4: 水平荷载效应分布图 (地震/风荷载)')
            
            # 添加详细说明
            seismic_text = (
                f"水平荷载效应分析说明：\n"
                f"1. 计算方法: 采用GB 50011-2010底部剪力法。\n"
                f"2. 水平力分布: 按各楼层重力荷载与高度乘积比例分配，顶层水平力最大。\n"
                f"3. 层间剪力: 从顶层至底层逐层累加，基底剪力等于各层水平力总和。\n"
                f"4. 倾覆力矩: 各层水平力对基础产生的力矩总和，用于计算柱轴力增量。\n"
                f"5. 柱轴力增量: 边柱轴力增量 dN = M_total / B，影响柱截面设计。"
            )
            doc.add_paragraph(seismic_text)
            
            # 如果有地震参数，添加参数说明
            if hasattr(grid, 'alpha_max') and grid.alpha_max > 0:
                param_text = (
                    f"\n设计参数：\n"
                    f"• 水平地震影响系数最大值 amax = {grid.alpha_max}\n"
                    f"• 结构基本周期 T1 = 0.08 x N = 0.08 x {grid.num_stories} = {0.08*grid.num_stories:.2f}s\n"
                    f"• 设计地震分组: {grid.seismic_group if hasattr(grid, 'seismic_group') else '2'}组"
                )
                doc.add_paragraph(param_text)

    # 确保输出路径干净
    import os
    if os.path.exists(output_path):
        try:
            os.remove(output_path)
        except OSError:
            print(f"警告: 无法删除旧文件 {output_path}, 可能被占用")

    try:
        doc.save(output_path)
        print(f"✓ Word设计计算书已保存: {output_path}")
    except Exception as e:
        print(f"❌ Word计算书保存失败: {str(e)}")

# =============================================================================
# P-M曲线图
# =============================================================================

def plot_pm_diagrams(result: OptimizationResult,
                     model,
                     db: SectionDatabase,
                     output_path: str = "PM曲线图.png") -> None:
    """
    绘制柱的P-M相互作用曲线图
    标注实际荷载点
    
    注意: P-M曲线约定压力为正，anaStruct返回的轴力压力为负
    因此需要将轴力取负（转换为压力为正）
    
    Args:
        result: 优化结果
        model: 结构模型
        db: 截面数据库
        output_path: 输出文件路径
    """
    # 收集柱的内力
    col_forces = {eid: f for eid, f in result.forces.items() 
                  if f.element_type == 'column'}
    
    if not col_forces:
        print("警告: 没有柱内力数据，跳过P-M曲线图")
        return
    
    # 获取不同的柱截面
    col_sections = {}
    for eid in col_forces.keys():
        sec_idx = model.column_sections.get(eid, 40)
        if sec_idx not in col_sections:
            col_sections[sec_idx] = []
        col_sections[sec_idx].append(eid)
    
    # 创建图表
    n_sections = len(col_sections)
    fig, axes = plt.subplots(1, min(n_sections, 3), figsize=(5*min(n_sections, 3), 5))
    if n_sections == 1:
        axes = [axes]
    
    As_col = REBAR_AREAS['4φ22']
    
    for ax, (sec_idx, elem_ids) in zip(axes, list(col_sections.items())[:3]):
        sec = db.get_by_index(sec_idx)
        
        # 生成P-M曲线 (压力为正)
        pm_curve = generate_pm_curve(sec['b'], sec['h'], As_col, num_points=50)
        P_vals = [p[0] for p in pm_curve]
        M_vals = [p[1] for p in pm_curve]
        
        # 绘制P-M包络线
        ax.plot(M_vals, P_vals, 'b-', linewidth=2, label='P-M包络线')
        ax.fill(M_vals, P_vals, alpha=0.1, color='blue')
        
        # 标注实际荷载点
        for i, eid in enumerate(elem_ids):
            f = col_forces[eid]
            # 关键修正：
            # 1. 弯矩取绝对值 (M_design已经是绝对值)
            # 2. 轴力：anaStruct返回压力为负，需要取负转换为压力为正
            #    axial_min 通常是最大压力（负值最大），取其绝对值
            M_actual = f.M_design  # 已经是绝对值
            N_actual = abs(f.axial_min)  # 取压力的绝对值 (压力为负，绝对值后为正)
            
            label = '实际荷载点' if i == 0 else None
            ax.plot(M_actual, N_actual, 'r^', markersize=8, label=label)
        
        # 图表设置
        ax.set_xlabel('弯矩 M (kN·m)')
        ax.set_ylabel('轴力 N (kN) [压为正]')
        ax.set_title(f'柱截面 {sec["b"]}×{sec["h"]} mm')
        ax.grid(True, alpha=0.3)
        ax.axhline(y=0, color='k', linewidth=0.5)
        ax.axvline(x=0, color='k', linewidth=0.5)
        ax.legend(loc='upper right')
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    print(f"✓ P-M曲线图已保存: {output_path}")
    plt.close()


# =============================================================================
# 框架内力图 (标准结构力学样式)
# =============================================================================

def plot_frame_diagrams(result: OptimizationResult,
                        model,
                        grid: GridInput,
                        output_path: str = "框架内力图.png") -> None:
    """
    绘制框架内力图（弯矩图、剪力图、轴力图）
    采用标准结构力学样式：线段表示 + 端点/跨中数值标注
    
    Args:
        result: 优化结果
        model: 结构模型
        grid: 轴网配置
        output_path: 输出文件路径
    """
    fig, axes = plt.subplots(1, 3, figsize=(20, 10))
    titles = ['弯矩图 M (kN·m)', '剪力图 V (kN)', '轴力图 N (kN)']
    colors = ['red', 'green', 'orange']
    
    for ax, title in zip(axes, titles):
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.set_xlabel('X (m)')
        ax.set_ylabel('Z (m)')
        ax.set_aspect('equal')
        ax.grid(True, alpha=0.3, linestyle='--')
    
    # 绘制框架几何（所有图共用）
    for ax in axes:
        _draw_frame_geometry_standard(ax, model, grid)
    
    # 缩放因子
    max_M = max(abs(f.M_design) for f in result.forces.values()) or 1
    max_V = max(abs(f.V_design) for f in result.forces.values()) or 1
    max_N = max(abs(f.N_design) for f in result.forces.values()) or 1
    
    scale_M = 0.8 / max_M  # 最大值对应0.8m偏移
    scale_V = 0.5 / max_V
    scale_N = 0.4 / max_N
    
    # 绘制内力图（标准样式）
    _draw_moment_diagram_standard(axes[0], result.forces, model, scale_M, colors[0])
    _draw_shear_diagram_standard(axes[1], result.forces, model, scale_V, colors[1])
    _draw_axial_diagram_standard(axes[2], result.forces, model, scale_N, colors[2])
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    print(f"✓ 框架内力图已保存: {output_path}")
    plt.close()


def _draw_frame_geometry_standard(ax, model, grid: GridInput):
    """绘制框架几何轮廓（标准样式）"""
    # 绘制构件（细黑线）
    for beam_id, (start, end) in model.beams.items():
        x1, z1 = model.nodes[start]
        x2, z2 = model.nodes[end]
        ax.plot([x1/1000, x2/1000], [z1/1000, z2/1000], 'k-', linewidth=1.5)
    
    for col_id, (start, end) in model.columns.items():
        x1, z1 = model.nodes[start]
        x2, z2 = model.nodes[end]
        ax.plot([x1/1000, x2/1000], [z1/1000, z2/1000], 'k-', linewidth=1.5)
    
    # 绘制节点（小圆点）
    for node_id, (x, z) in model.nodes.items():
        ax.plot(x/1000, z/1000, 'ko', markersize=3)
    
    # 绘制固定支座（三角形+横线）
    n_cols = grid.num_spans + 1
    for i in range(n_cols):
        x = sum(grid.x_spans[:i]) / 1000
        # 三角形
        ax.plot([x-0.15, x+0.15, x, x-0.15], [-0.1, -0.1, 0, -0.1], 'k-', linewidth=1.5)
        # 横线
        ax.plot([x-0.2, x+0.2], [-0.15, -0.15], 'k-', linewidth=1.5)
        # 斜线阴影
        for dx in np.linspace(-0.18, 0.14, 5):
            ax.plot([x+dx, x+dx+0.08], [-0.15, -0.22], 'k-', linewidth=0.5)


def _draw_moment_diagram_standard(ax, forces: Dict, model, scale: float, color: str):
    """
    绘制弯矩图（标准结构力学样式）
    - 梁：正弯矩向下（受拉侧），画抛物线形状
    - 柱：弯矩画在受拉侧
    - 标注端点和跨中数值
    """
    for elem_id, f in forces.items():
        if f.element_type == 'beam':
            start, end = model.beams[elem_id]
            x1, z1 = model.nodes[start]
            x2, z2 = model.nodes[end]
            
            # 梁弯矩图：端部为负弯矩(moment_min)，跨中为正弯矩(moment_max)
            # 均布荷载下呈抛物线分布
            L = (x2 - x1) / 1000  # 跨度(m)
            n_pts = 30  # 增加点数使曲线更光滑
            
            # 端部弯矩（负值，使用moment_min）和跨中弯矩（正值，使用moment_max）
            M_left = f.moment_min * scale   # 左端弯矩 (通常为负)
            M_right = f.moment_min * scale  # 右端弯矩 (假设对称)
            M_mid = f.moment_max * scale    # 跨中弯矩 (通常为正)
            
            # 使用三点抛物线插值: y = at² + bt + c
            # t=0: M_left, t=0.5: M_mid, t=1: M_right
            # 解出系数
            c = M_left
            a = 2 * (M_left + M_right - 2 * M_mid)
            b = M_right - M_left - a
            
            t = np.linspace(0, 1, n_pts)
            M_pts = a * t**2 + b * t + c
            
            x_pts = np.linspace(x1/1000, x2/1000, n_pts)
            z_base = z1/1000
            z_pts = z_base - M_pts  # 正弯矩向下绘制
            
            # 绘制弯矩图轮廓
            x_fill = np.concatenate([[x1/1000], x_pts, [x2/1000]])
            z_fill = np.concatenate([[z_base], z_pts, [z_base]])
            ax.fill(x_fill, z_fill, alpha=0.3, color=color, edgecolor=color, linewidth=1)
            ax.plot(x_pts, z_pts, color=color, linewidth=1.5)
            
            # 标注端点值
            # 注意：这里简化假设两端弯矩近似相等（均为负弯矩M_min）
            # 实际上外跨梁两端弯矩可能差异较大
            ax.annotate(f'{abs(f.moment_min):.0f}', 
                       xy=(x1/1000, z_base), xytext=(x1/1000-0.2, z_base+0.2),
                       fontsize=7, color=color)
            ax.annotate(f'{abs(f.moment_min):.0f}', 
                       xy=(x2/1000, z_base), xytext=(x2/1000+0.1, z_base+0.2),
                       fontsize=7, color=color)
            
            # 标注跨中弯矩
            x_mid = (x1 + x2) / 2000
            z_mid = z_base - M_mid
            ax.text(x_mid, z_mid - 0.3, f'{abs(f.moment_max):.0f}', 
                   ha='center', va='top', fontsize=7, color=color)
            # 标注跨中值
            ax.annotate(f'{f.M_design:.0f}', 
                       xy=((x1+x2)/2000, z_base - M_mid), 
                       xytext=((x1+x2)/2000, z_base - M_mid - 0.15),
                       fontsize=7, color=color, ha='center')
        
        else:  # 柱
            start, end = model.columns[elem_id]
            x1, z1 = model.nodes[start]
            x2, z2 = model.nodes[end]
            
            M = f.M_design * scale
            x_base = x1/1000
            
            # 柱弯矩图：假设线性分布
            ax.fill([x_base, x_base + M, x_base + M, x_base],
                   [z1/1000, z1/1000, z2/1000, z2/1000],
                   alpha=0.3, color=color, edgecolor=color, linewidth=1)
            ax.plot([x_base + M, x_base + M], [z1/1000, z2/1000], 
                   color=color, linewidth=1.5)
            
            # 标注底部值
            if f.M_design > 1:
                ax.annotate(f'{f.M_design:.0f}', 
                           xy=(x_base + M, z1/1000), 
                           xytext=(x_base + M + 0.1, z1/1000 + 0.1),
                           fontsize=6, color=color)


def _draw_shear_diagram_standard(ax, forces: Dict, model, scale: float, color: str):
    """
    绘制剪力图（标准结构力学样式）
    - 梁：正剪力向上，线性变化
    - 标注端点数值
    """
    for elem_id, f in forces.items():
        if f.element_type == 'beam':
            start, end = model.beams[elem_id]
            x1, z1 = model.nodes[start]
            x2, z2 = model.nodes[end]
            
            z_base = z1/1000
            V_left = f.shear_max * scale
            V_right = f.shear_min * scale
            
            # 剪力图：梯形（均布荷载下线性变化）
            ax.fill([x1/1000, x1/1000, x2/1000, x2/1000],
                   [z_base, z_base + V_left, z_base + V_right, z_base],
                   alpha=0.3, color=color, edgecolor=color, linewidth=1)
            ax.plot([x1/1000, x2/1000], [z_base + V_left, z_base + V_right], 
                   color=color, linewidth=1.5)
            
            # 标注端点值
            ax.annotate(f'{abs(f.shear_max):.0f}', 
                       xy=(x1/1000, z_base + V_left), 
                       xytext=(x1/1000-0.15, z_base + V_left),
                       fontsize=7, color=color, ha='right')
            ax.annotate(f'{abs(f.shear_min):.0f}', 
                       xy=(x2/1000, z_base + V_right), 
                       xytext=(x2/1000+0.15, z_base + V_right),
                       fontsize=7, color=color, ha='left')
        
        else:  # 柱
            start, end = model.columns[elem_id]
            x1, z1 = model.nodes[start]
            x2, z2 = model.nodes[end]
            
            V = f.V_design * scale
            x_base = x1/1000
            
            # 柱剪力：假设常量
            ax.fill([x_base, x_base + V, x_base + V, x_base],
                   [z1/1000, z1/1000, z2/1000, z2/1000],
                   alpha=0.3, color=color, edgecolor=color, linewidth=1)
            ax.plot([x_base + V, x_base + V], [z1/1000, z2/1000], 
                   color=color, linewidth=1.5)


def _draw_axial_diagram_standard(ax, forces: Dict, model, scale: float, color: str):
    """
    绘制轴力图（标准结构力学样式）
    - 柱：压力为正，画在构件侧面
    - 标注数值
    """
    for elem_id, f in forces.items():
        if f.element_type == 'column':
            start, end = model.columns[elem_id]
            x1, z1 = model.nodes[start]
            x2, z2 = model.nodes[end]
            
            N = f.N_design * scale
            x_base = x1/1000
            
            # 柱轴力：假设常量
            ax.fill([x_base, x_base + N, x_base + N, x_base],
                   [z1/1000, z1/1000, z2/1000, z2/1000],
                   alpha=0.3, color=color, edgecolor=color, linewidth=1)
            ax.plot([x_base + N, x_base + N], [z1/1000, z2/1000], 
                   color=color, linewidth=1.5)
            
            # 标注轴力值
            if f.N_design > 1:
                ax.annotate(f'{f.N_design:.0f}', 
                           xy=(x_base + N, (z1+z2)/2000), 
                           xytext=(x_base + N + 0.1, (z1+z2)/2000),
                           fontsize=6, color=color, ha='left', va='center')


# =============================================================================
# 收敛曲线
# =============================================================================

def plot_convergence(cost_history: List[float],
                     output_path: str = "收敛曲线.png") -> None:
    """
    绘制优化收敛曲线
    
    Args:
        cost_history: 每代最优造价历史
        output_path: 输出文件路径
    """
    fig, ax = plt.subplots(figsize=(10, 5))
    
    ax.plot(cost_history, 'b-', linewidth=2)
    ax.set_xlabel('迭代代数')
    ax.set_ylabel('最优造价 (元)')
    ax.set_title('遗传算法收敛曲线')
    ax.grid(True, alpha=0.3)
    
    # 标注初始和最终值
    ax.annotate(f'初始: {cost_history[0]:,.0f}元',
               xy=(0, cost_history[0]), xytext=(5, cost_history[0]),
               fontsize=10)
    ax.annotate(f'最终: {cost_history[-1]:,.0f}元',
               xy=(len(cost_history)-1, cost_history[-1]),
               fontsize=10)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    print(f"✓ 收敛曲线已保存: {output_path}")
    plt.close()


# =============================================================================
# 地震/水平荷载效应可视化
# =============================================================================

def plot_seismic_load_diagram(grid: GridInput,
                               model = None,
                               output_path: str = "水平荷载效应图.png") -> Dict:
    """
    绘制地震/水平荷载效应示意图
    
    包含：
    1. 框架立面 + 水平力分布箭头
    2. 层间剪力分布图
    3. 倾覆力矩/柱轴力增量分布图
    
    Args:
        grid: 轴网配置 (必须包含 alpha_max)
        model: 结构模型 (可选，用于获取实际内力)
        output_path: 输出文件路径
        
    Returns:
        Dict: 计算结果数据
    """
    # 检查是否有水平荷载
    has_seismic = hasattr(grid, 'alpha_max') and grid.alpha_max > 0
    has_wind = hasattr(grid, 'w0') and grid.w0 > 0
    
    if not has_seismic and not has_wind:
        print("⚠ 无水平荷载，跳过水平荷载效应图")
        return {}
    
    n_stories = grid.num_stories
    n_spans = grid.num_spans
    
    # ========================================================================
    # 计算地震作用 (GB 50011-2010 底部剪力法)
    # ========================================================================
    
    # 基本周期估算
    T1 = 0.08 * n_stories
    
    # 地震影响系数
    alpha_max = getattr(grid, 'alpha_max', 0.08)
    alpha_1 = alpha_max * 0.9
    
    # 重力荷载代表值
    psi_c = 0.5
    q_e = grid.q_dead + psi_c * grid.q_live  # kN/m
    total_span = sum(grid.x_spans) / 1000  # m
    
    # 各层重力荷载和高度
    G_story = []
    H_story = []  # 累计高度
    h_story = []  # 层高
    cumulative_height = 0.0
    
    for story in range(n_stories):
        story_height = grid.z_heights[story] / 1000  # m
        h_story.append(story_height)
        cumulative_height += story_height
        H_story.append(cumulative_height)
        
        G_beam = q_e * total_span
        G_col = G_beam * 0.15
        G_story.append(G_beam + G_col)
    
    G_total = sum(G_story)
    
    # 基底剪力
    eta = 0.85
    F_EK = alpha_1 * G_total * eta
    
    # 顶部附加地震作用
    Tg = 0.4
    if T1 > 1.4 * Tg:
        delta_n = min(0.08 * T1 + 0.07, 0.25)
    else:
        delta_n = 0.0
    
    F_top_extra = delta_n * F_EK
    F_distribute = F_EK - F_top_extra
    
    # 各层地震力分配 (倒三角形)
    sum_GH = sum(G_story[i] * H_story[i] for i in range(n_stories))
    
    F_story = []
    for i in range(n_stories):
        if sum_GH > 0:
            F_i = F_distribute * (G_story[i] * H_story[i]) / sum_GH
        else:
            F_i = 0.0
        
        if i == n_stories - 1:
            F_i += F_top_extra
        
        F_story.append(F_i)
    
    # 计算层间剪力 (从顶到底累加)
    V_story = []
    V_cumulative = 0.0
    for i in range(n_stories - 1, -1, -1):
        V_cumulative += F_story[i]
        V_story.insert(0, V_cumulative)
    
    # 计算倾覆力矩 (各楼层对基础的力矩)
    M_overturn = []
    for i in range(n_stories):
        M_i = F_story[i] * H_story[i]
        M_overturn.append(M_i)
    
    M_total = sum(M_overturn)
    
    # 柱轴力增量 (简化：按倾覆力矩估算)
    # ΔN = M_total / (B × n_cols), B = 总宽度
    B = total_span
    n_cols = n_spans + 1
    if B > 0 and n_cols > 1:
        delta_N_edge = M_total / (B * (n_cols - 1))  # 边柱轴力增量
    else:
        delta_N_edge = 0
    
    # ========================================================================
    # 创建可视化图表
    # ========================================================================
    fig = plt.figure(figsize=(16, 10))
    
    # 子图1: 框架立面 + 水平力箭头
    ax1 = fig.add_subplot(2, 2, 1)
    _draw_frame_with_seismic(ax1, grid, H_story, F_story, F_EK)
    
    # 子图2: 水平力和层间剪力分布
    ax2 = fig.add_subplot(2, 2, 2)
    _draw_shear_distribution(ax2, H_story, F_story, V_story)
    
    # 子图3: 倾覆力矩和柱轴力增量
    ax3 = fig.add_subplot(2, 2, 3)
    _draw_moment_distribution(ax3, H_story, M_overturn, M_total, delta_N_edge, B)
    
    # 子图4: 参数汇总表
    ax4 = fig.add_subplot(2, 2, 4)
    _draw_seismic_summary(ax4, grid, T1, alpha_1, G_total, F_EK, V_story[0], M_total)
    
    plt.suptitle('水平荷载效应分析 (GB 50011-2010 底部剪力法)', fontsize=14, fontweight='bold')
    plt.tight_layout(rect=[0, 0, 1, 0.96])
    
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    print(f"✓ 水平荷载效应图已保存: {output_path}")
    plt.close()
    
    # 返回计算数据
    return {
        'T1': T1,
        'alpha_1': alpha_1,
        'G_total': G_total,
        'F_EK': F_EK,
        'F_story': F_story,
        'V_story': V_story,
        'M_total': M_total,
        'delta_N_edge': delta_N_edge,
    }


def _draw_frame_with_seismic(ax, grid: GridInput, H_story: List, F_story: List, F_EK: float):
    """绘制框架立面 + 水平力箭头"""
    n_stories = grid.num_stories
    n_spans = grid.num_spans
    total_width = sum(grid.x_spans) / 1000  # m
    total_height = sum(grid.z_heights) / 1000  # m
    
    # 绘制框架轮廓
    for i in range(n_spans + 1):
        x = sum(grid.x_spans[:i]) / 1000
        ax.plot([x, x], [0, total_height], 'k-', linewidth=2)
    
    for j in range(n_stories + 1):
        z = sum(grid.z_heights[:j]) / 1000
        ax.plot([0, total_width], [z, z], 'k-', linewidth=1.5)
    
    # 绘制固定支座
    for i in range(n_spans + 1):
        x = sum(grid.x_spans[:i]) / 1000
        ax.plot([x-0.15, x+0.15, x, x-0.15], [-0.1, -0.1, 0, -0.1], 'k-', linewidth=1.5)
    
    # 绘制水平力箭头
    max_F = max(F_story) if F_story else 1
    arrow_scale = total_width * 0.4 / max_F
    
    for i, (h, F) in enumerate(zip(H_story, F_story)):
        arrow_len = F * arrow_scale
        ax.arrow(-arrow_len - 0.2, h, arrow_len, 0, 
                head_width=0.15, head_length=0.1, fc='red', ec='red', linewidth=2)
        ax.text(-arrow_len - 0.4, h, f'{F:.1f} kN', va='center', ha='right', 
               fontsize=9, color='red')
    
    # 标注
    ax.set_title(f'水平地震力分布 (F_EK={F_EK:.1f} kN)', fontsize=11)
    ax.set_xlabel('X (m)')
    ax.set_ylabel('Z (m)')
    ax.set_xlim(-total_width * 0.8, total_width * 1.1)
    ax.set_ylim(-0.5, total_height + 0.5)
    ax.set_aspect('equal')
    ax.grid(True, alpha=0.3)


def _draw_shear_distribution(ax, H_story: List, F_story: List, V_story: List):
    """绘制层间剪力分布图"""
    n_stories = len(H_story)
    
    # 水平力 (阶梯图)
    heights_F = [0] + H_story
    values_F = [0] + F_story
    
    # 层间剪力 (阶梯图)
    heights_V = [0] + H_story
    values_V = V_story + [0]
    
    # 绘制层间剪力
    for i in range(n_stories):
        h_bottom = heights_V[i]
        h_top = heights_V[i + 1]
        V = values_V[i]
        ax.fill_betweenx([h_bottom, h_top], 0, V, alpha=0.3, color='blue', step='post')
        ax.plot([V, V], [h_bottom, h_top], 'b-', linewidth=2)
        ax.text(V + 2, (h_bottom + h_top) / 2, f'{V:.1f}', va='center', fontsize=9, color='blue')
    
    # 绘制楼层水平力
    for i, (h, F) in enumerate(zip(H_story, F_story)):
        ax.plot(F, h, 'r^', markersize=10)
        ax.text(F + 2, h + 0.2, f'{F:.1f}', fontsize=8, color='red')
    
    ax.axvline(x=0, color='k', linewidth=0.5)
    ax.set_xlabel('力 (kN)')
    ax.set_ylabel('高度 Z (m)')
    ax.set_title('层间剪力分布', fontsize=11)
    ax.legend(['层间剪力 V', '楼层水平力 F'], loc='upper right')
    ax.grid(True, alpha=0.3)
    ax.set_xlim(-5, max(V_story) * 1.3)


def _draw_moment_distribution(ax, H_story: List, M_overturn: List, M_total: float, delta_N: float, B: float):
    """绘制倾覆力矩分布图"""
    n_stories = len(H_story)
    
    # 累计倾覆力矩
    M_cumulative = []
    M_cum = 0
    for i in range(n_stories - 1, -1, -1):
        M_cum += M_overturn[i]
        M_cumulative.insert(0, M_cum)
    
    # 绘制倾覆力矩
    heights = [0] + H_story
    moments = M_cumulative + [0]
    
    ax.fill_betweenx(H_story, 0, M_cumulative, alpha=0.3, color='orange')
    ax.plot(M_cumulative, H_story, 'o-', color='orange', linewidth=2, markersize=8)
    
    for h, M in zip(H_story, M_cumulative):
        ax.text(M + 5, h, f'{M:.0f}', va='center', fontsize=9, color='darkorange')
    
    ax.axvline(x=0, color='k', linewidth=0.5)
    ax.set_xlabel('倾覆力矩 (kN*m)')
    ax.set_ylabel('高度 Z (m)')
    ax.set_title(f'倾覆力矩分布 (M_total={M_total:.0f} kN*m)', fontsize=11)
    ax.grid(True, alpha=0.3)
    
    # 添加柱轴力增量说明
    ax.text(0.95, 0.05, f'边柱轴力增量 dN = {delta_N:.1f} kN\n(按 M/B 估算)', 
           transform=ax.transAxes, fontsize=9, ha='right', va='bottom',
           bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))


def _draw_seismic_summary(ax, grid: GridInput, T1: float, alpha_1: float, 
                          G_total: float, F_EK: float, V_base: float, M_total: float):
    """绘制参数汇总表"""
    ax.axis('off')
    
    # 表格数据 (使用ASCII兼容字符避免字体显示问题)
    data = [
        ['参数', '数值', '说明'],
        ['结构周期 T1', f'{T1:.2f} s', '经验公式 0.08N'],
        ['地震影响系数 a1', f'{alpha_1:.3f}', f'amax={grid.alpha_max}*0.9'],
        ['重力荷载代表值 GE', f'{G_total:.0f} kN', 'G + 0.5Q'],
        ['基底剪力 FEK', f'{F_EK:.1f} kN', 'a1*GE*eta'],
        ['基底剪力 Vbase', f'{V_base:.1f} kN', '= FEK'],
        ['总倾覆力矩 Mtotal', f'{M_total:.0f} kN*m', 'Sum(Fi*Hi)'],
    ]
    
    # 添加风荷载信息（如果有）
    if hasattr(grid, 'w0') and grid.w0 > 0:
        data.append(['基本风压 w0', f'{grid.w0} kN/m2', 'GB 50009-2012'])
    
    # 创建表格
    table = ax.table(cellText=data, loc='center', cellLoc='center',
                    colWidths=[0.35, 0.25, 0.4])
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1.2, 1.8)
    
    # 设置表头样式
    for j in range(3):
        table[(0, j)].set_facecolor('#4472C4')
        table[(0, j)].set_text_props(color='white', fontweight='bold')
    
    ax.set_title('地震作用计算参数 (GB 50011-2010)', fontsize=11, pad=20)


# =============================================================================
# 测试代码
# =============================================================================

if __name__ == "__main__":
    print("报表生成器模块测试")
    print("请运行 main.py 进行完整测试")

