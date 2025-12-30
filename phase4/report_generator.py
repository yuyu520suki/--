"""
报表生成器模块 - Excel报表和可视化图表
包括P-M曲线图、框架内力图、收敛曲线
"""

import sys
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
from datetime import datetime

# 设置中文字体
matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False

# 添加父目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from phase1.section_database import SectionDatabase
from phase1.capacity_calculator import generate_pm_curve, REBAR_AREAS
from phase4.data_models import GridInput, ElementForces, OptimizationResult


# =============================================================================
# Excel 报表生成
# =============================================================================

def generate_excel_report(result: OptimizationResult,
                          model,
                          db: SectionDatabase,
                          output_path: str = "优化结果.xlsx") -> None:
    """
    生成Excel优化结果报表
    
    Args:
        result: 优化结果
        model: 结构模型
        db: 截面数据库
        output_path: 输出文件路径
    """
    try:
        import pandas as pd
    except ImportError:
        print("警告: pandas未安装，跳过Excel报表生成")
        return
    
    # 构建数据
    rows = []
    
    for elem_id, forces in result.forces.items():
        if forces.element_type == 'beam':
            sec_idx = model.beam_sections.get(elem_id, 30)
        else:
            sec_idx = model.column_sections.get(elem_id, 40)
        
        sec = db.get_by_index(sec_idx)
        
        # 计算利用率 (简化)
        from phase1.capacity_calculator import calculate_capacity
        As = REBAR_AREAS['3φ20'] if forces.element_type == 'beam' else REBAR_AREAS['4φ22']
        cap = calculate_capacity(sec['b'], sec['h'], As)
        
        utility_M = forces.M_design / cap['phi_Mn'] if cap['phi_Mn'] > 0 else 999
        utility_V = forces.V_design / cap['phi_Vn'] if cap['phi_Vn'] > 0 else 999
        
        # 配筋率
        rho = As / (sec['b'] * sec['h']) * 100
        
        rows.append({
            '单元ID': elem_id,
            '类型': '梁' if forces.element_type == 'beam' else '柱',
            '截面': f"{sec['b']}×{sec['h']}",
            '宽度b(mm)': sec['b'],
            '高度h(mm)': sec['h'],
            '配筋As(mm²)': As,
            '配筋率(%)': round(rho, 2),
            'M设计(kN·m)': round(forces.M_design, 1),
            'V设计(kN)': round(forces.V_design, 1),
            'N设计(kN)': round(forces.N_design, 1),
            'M利用率': round(utility_M, 3),
            'V利用率': round(utility_V, 3),
            '状态': '✓' if max(utility_M, utility_V) <= 1.0 else '✗',
        })
    
    df = pd.DataFrame(rows)
    
    # 按类型分组
    df_beams = df[df['类型'] == '梁'].copy()
    df_columns = df[df['类型'] == '柱'].copy()
    
    # 写入Excel
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        # 汇总表
        summary = pd.DataFrame({
            '项目': ['总造价', '梁数量', '柱数量', '最大梁M利用率', '最大柱M利用率'],
            '数值': [
                f"{result.cost:,.0f} 元",
                len(df_beams),
                len(df_columns),
                f"{df_beams['M利用率'].max():.2f}" if len(df_beams) > 0 else 'N/A',
                f"{df_columns['M利用率'].max():.2f}" if len(df_columns) > 0 else 'N/A',
            ]
        })
        summary.to_excel(writer, sheet_name='汇总', index=False)
        
        # 梁详细表
        df_beams.to_excel(writer, sheet_name='梁详细', index=False)
        
        # 柱详细表
        df_columns.to_excel(writer, sheet_name='柱详细', index=False)
    
    print(f"✓ Excel报表已保存: {output_path}")


# =============================================================================
# Word 计算书生成
# =============================================================================

def generate_word_report(result: OptimizationResult,
                         model,
                         db: SectionDatabase,
                         grid: GridInput,
                         output_path: str = "设计计算书.docx",
                         image_paths: Dict[str, str] = None) -> None:
    """
    生成详细的Word设计计算书
    
    Args:
        result: 优化结果
        model: 结构模型
        db: 截面数据库
        grid: 轴网信息
        output_path: 输出路径
        image_paths: 图片路径字典 {'pm': path, 'frame': path, 'conv': path}
    """
    try:
        from docx import Document
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("警告: python-docx未安装，跳过Word计算书生成")
        return

    doc = Document()
    
    # --- 标题 ---
    heading = doc.add_heading('RC框架结构优化设计计算书', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('Design by Antigravity AI Optimization System')
    
    # --- 1. 工程概况 ---
    doc.add_heading('1. 工程概况', level=1)
    p = doc.add_paragraph()
    p.add_run(f'本工程为{grid.num_stories}层钢筋混凝土框架结构，').bold = False
    p.add_run(f'总高度 {grid.total_height/1000:.1f}m，总宽度 {grid.total_width/1000:.1f}m。').bold = False
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table.cell(0, 0).text = '结构形式'
    table.cell(0, 1).text = '现浇钢筋混凝土框架结构'
    table.cell(1, 0).text = '跨度信息'
    table.cell(1, 1).text = f'{grid.num_spans}跨 ({", ".join([str(x/1000)+"m" for x in grid.x_spans])})'
    table.cell(2, 0).text = '层高信息'
    table.cell(2, 1).text = f'首层 {grid.z_heights[0]/1000}m, 标准层 {grid.z_heights[1]/1000}m'
    table.cell(3, 0).text = '结构重要性系数'
    table.cell(3, 1).text = '1.0 (丙类建筑)'

    # --- 2. 设计依据与计算方法 ---
    doc.add_heading('2. 设计依据与方法', level=1)
    
    doc.add_heading('2.1 执行规范', level=2)
    doc.add_paragraph('1. 《混凝土结构设计规范》 (GB 50010-2010)')
    doc.add_paragraph('2. 《建筑结构荷载规范》 (GB 50009-2012)')
    
    doc.add_heading('2.2 荷载信息', level=2)
    doc.add_paragraph(f'• 恒载 (Dead Load): {grid.q_dead} kN/m')
    doc.add_paragraph(f'• 活载 (Live Load): {grid.q_live} kN/m')
    if hasattr(grid, 'w0') and grid.w0 > 0:
        doc.add_paragraph(f'• 基本风压: {grid.w0} kN/m² (GB 50009-2012)')
    if hasattr(grid, 's0') and grid.s0 > 0:
        doc.add_paragraph(f'• 基本雪压: {grid.s0} kN/m² (GB 50009-2012)')
    
    doc.add_heading('2.3 荷载组合 (GB 50009-2012)', level=2)
    combo_table = doc.add_table(rows=1, cols=2)
    combo_table.style = 'Table Grid'
    combo_hdr = combo_table.rows[0].cells
    combo_hdr[0].text = '极限状态'
    combo_hdr[1].text = '荷载组合'
    
    combos = [
        ('承载能力 ULS', '1.2G+1.4Q, 1.35G+0.98Q'),
        ('正常使用 SLS', 'G+Q (标准), G+0.4Q (准永久)'),
    ]
    if hasattr(grid, 'w0') and grid.w0 > 0:
        combos[0] = ('承载能力 ULS', '1.2G+1.4Q, 1.2G+1.4W, 1.2G+0.98Q+1.4W')
    if hasattr(grid, 's0') and grid.s0 > 0:
        combos[0] = ('承载能力 ULS', combos[0][1] + ', 1.2G+1.4S')
    
    for state, combo in combos:
        row = combo_table.add_row().cells
        row[0].text = state
        row[1].text = combo
    
    doc.add_heading('2.4 验算项目 (GB 50010-2010)', level=2)
    doc.add_paragraph('• 正截面承载力验算 (受弯、偏压)')
    doc.add_paragraph('• 斜截面承载力验算 (受剪)')
    doc.add_paragraph('• 柱轴压比限值验算 (μ ≤ 0.9)')
    doc.add_paragraph('• 最小/最大配筋率检查')
    doc.add_paragraph('• 挠度限值验算 (l/200~l/250)')
    doc.add_paragraph('• 裂缝宽度限值验算 (≤0.3mm)')
    
    doc.add_heading('2.5 计算方法', level=2)
    doc.add_paragraph('• 内力分析: 采用矩阵位移法 (anaStruct有限元内核)')
    doc.add_paragraph('• 截面设计: 采用遗传算法 (PyGAD) 进行全局优化')
    doc.add_paragraph('• 承载力验算: 考虑P-M相互作用 (柱) 和双向受力 (梁)')

    # --- 3. 结构选型结果 ---
    doc.add_heading('3. 优化结果', level=1)
    doc.add_paragraph(f'经过 {len(result.fitness_history)} 代遗传进化，得到最优设计方案如下：')
    doc.add_paragraph(f'总造价: ¥{result.cost:,.2f}').runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    # 截面配置表
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '构件类型'
    hdr_cells[1].text = '截面尺寸 (mm)'
    hdr_cells[2].text = '配筋配置'
    hdr_cells[3].text = '配筋率'
    
    group_map = [
        ('标准层梁', 0, 'beam', '3φ20'),
        ('屋面梁', 1, 'beam', '3φ20'),
        ('角柱', 2, 'column', '4φ22'),
        ('内柱', 3, 'column', '4φ22')
    ]
    
    for name, idx, type_, rebar in group_map:
        if idx < len(result.genes):
            sec_idx = result.genes[idx]
            sec = db.get_by_index(sec_idx)
            row = table.add_row().cells
            row[0].text = name
            row[1].text = f"{sec['b']} × {sec['h']}"
            row[2].text = rebar
            
            As = REBAR_AREAS[rebar]
            rho = As / (sec['b'] * sec['h']) * 100
            row[3].text = f"{rho:.2f}%"

    # --- 4. 附图 ---
    doc.add_heading('4. 计算附图', level=1)
    
    if image_paths:
        if 'frame' in image_paths and Path(image_paths['frame']).exists():
            doc.add_heading('4.1 框架内力图', level=2)
            doc.add_picture(image_paths['frame'], width=Cm(16))
            doc.add_paragraph('图 4-1: 框架弯矩、剪力、轴力包络图')
            
        if 'pm' in image_paths and Path(image_paths['pm']).exists():
            doc.add_heading('4.2 柱P-M相互作用曲线', level=2)
            doc.add_picture(image_paths['pm'], width=Cm(16))
            doc.add_paragraph('图 4-2: 柱截面安全包络及荷载校核')
            
        if 'conv' in image_paths and Path(image_paths['conv']).exists():
            doc.add_heading('4.3 优化收敛过程', level=2)
            doc.add_picture(image_paths['conv'], width=Cm(12))
            doc.add_paragraph('图 4-3: 遗传算法造价收敛曲线')
            
            # 动态计算收敛代数 (变化率小于1%时认为收敛)
            if result.convergence_history and len(result.convergence_history) > 5:
                history = result.convergence_history
                final_cost = history[-1]
                convergence_gen = len(history)
                for i in range(len(history) - 1):
                    if abs(history[i] - final_cost) / final_cost < 0.01:
                        convergence_gen = i + 1
                        break
                doc.add_paragraph(f'收敛代数: 约第 {convergence_gen} 代进入稳定期 (变化率<1%)')
            else:
                doc.add_paragraph(f'总迭代代数: {len(result.fitness_history)} 代')

    doc.save(output_path)
    print(f"✓ Word设计计算书已保存: {output_path}")

# =============================================================================
# P-M曲线图
# =============================================================================

def plot_pm_diagrams(result: OptimizationResult,
                     model,
                     db: SectionDatabase,
                     output_path: str = "PM曲线图.png") -> None:
    """
    绘制柱的P-M相互作用曲线图
    标注实际荷载点
    
    Args:
        result: 优化结果
        model: 结构模型
        db: 截面数据库
        output_path: 输出文件路径
    """
    # 收集柱的内力
    col_forces = {eid: f for eid, f in result.forces.items() 
                  if f.element_type == 'column'}
    
    if not col_forces:
        print("警告: 没有柱内力数据，跳过P-M曲线图")
        return
    
    # 获取不同的柱截面
    col_sections = {}
    for eid in col_forces.keys():
        sec_idx = model.column_sections.get(eid, 40)
        if sec_idx not in col_sections:
            col_sections[sec_idx] = []
        col_sections[sec_idx].append(eid)
    
    # 创建图表
    n_sections = len(col_sections)
    fig, axes = plt.subplots(1, min(n_sections, 3), figsize=(5*min(n_sections, 3), 5))
    if n_sections == 1:
        axes = [axes]
    
    As_col = REBAR_AREAS['4φ22']
    
    for ax, (sec_idx, elem_ids) in zip(axes, list(col_sections.items())[:3]):
        sec = db.get_by_index(sec_idx)
        
        # 生成P-M曲线
        pm_curve = generate_pm_curve(sec['b'], sec['h'], As_col, num_points=50)
        P_vals = [p[0] for p in pm_curve]
        M_vals = [p[1] for p in pm_curve]
        
        # 绘制P-M包络线
        ax.plot(M_vals, P_vals, 'b-', linewidth=2, label='P-M包络线')
        ax.fill(M_vals, P_vals, alpha=0.1, color='blue')
        
        # 标注实际荷载点
        for i, eid in enumerate(elem_ids):
            f = col_forces[eid]
            label = '实际荷载点' if i == 0 else None
            ax.plot(abs(f.M_design), f.N_design, 'r^', markersize=8, label=label)
        
        # 图表设置
        ax.set_xlabel('弯矩 M (kN·m)')
        ax.set_ylabel('轴力 N (kN)')
        ax.set_title(f'柱截面 {sec["b"]}×{sec["h"]} mm')
        ax.grid(True, alpha=0.3)
        ax.axhline(y=0, color='k', linewidth=0.5)
        ax.axvline(x=0, color='k', linewidth=0.5)
        ax.legend(loc='upper right')
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    print(f"✓ P-M曲线图已保存: {output_path}")
    plt.close()


# =============================================================================
# 框架内力图 (标准结构力学样式)
# =============================================================================

def plot_frame_diagrams(result: OptimizationResult,
                        model,
                        grid: GridInput,
                        output_path: str = "框架内力图.png") -> None:
    """
    绘制框架内力图（弯矩图、剪力图、轴力图）
    采用标准结构力学样式：线段表示 + 端点/跨中数值标注
    
    Args:
        result: 优化结果
        model: 结构模型
        grid: 轴网配置
        output_path: 输出文件路径
    """
    fig, axes = plt.subplots(1, 3, figsize=(20, 10))
    titles = ['弯矩图 M (kN·m)', '剪力图 V (kN)', '轴力图 N (kN)']
    colors = ['red', 'green', 'orange']
    
    for ax, title in zip(axes, titles):
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.set_xlabel('X (m)')
        ax.set_ylabel('Z (m)')
        ax.set_aspect('equal')
        ax.grid(True, alpha=0.3, linestyle='--')
    
    # 绘制框架几何（所有图共用）
    for ax in axes:
        _draw_frame_geometry_standard(ax, model, grid)
    
    # 缩放因子
    max_M = max(abs(f.M_design) for f in result.forces.values()) or 1
    max_V = max(abs(f.V_design) for f in result.forces.values()) or 1
    max_N = max(abs(f.N_design) for f in result.forces.values()) or 1
    
    scale_M = 0.8 / max_M  # 最大值对应0.8m偏移
    scale_V = 0.5 / max_V
    scale_N = 0.4 / max_N
    
    # 绘制内力图（标准样式）
    _draw_moment_diagram_standard(axes[0], result.forces, model, scale_M, colors[0])
    _draw_shear_diagram_standard(axes[1], result.forces, model, scale_V, colors[1])
    _draw_axial_diagram_standard(axes[2], result.forces, model, scale_N, colors[2])
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    print(f"✓ 框架内力图已保存: {output_path}")
    plt.close()


def _draw_frame_geometry_standard(ax, model, grid: GridInput):
    """绘制框架几何轮廓（标准样式）"""
    # 绘制构件（细黑线）
    for beam_id, (start, end) in model.beams.items():
        x1, z1 = model.nodes[start]
        x2, z2 = model.nodes[end]
        ax.plot([x1/1000, x2/1000], [z1/1000, z2/1000], 'k-', linewidth=1.5)
    
    for col_id, (start, end) in model.columns.items():
        x1, z1 = model.nodes[start]
        x2, z2 = model.nodes[end]
        ax.plot([x1/1000, x2/1000], [z1/1000, z2/1000], 'k-', linewidth=1.5)
    
    # 绘制节点（小圆点）
    for node_id, (x, z) in model.nodes.items():
        ax.plot(x/1000, z/1000, 'ko', markersize=3)
    
    # 绘制固定支座（三角形+横线）
    n_cols = grid.num_spans + 1
    for i in range(n_cols):
        x = sum(grid.x_spans[:i]) / 1000
        # 三角形
        ax.plot([x-0.15, x+0.15, x, x-0.15], [-0.1, -0.1, 0, -0.1], 'k-', linewidth=1.5)
        # 横线
        ax.plot([x-0.2, x+0.2], [-0.15, -0.15], 'k-', linewidth=1.5)
        # 斜线阴影
        for dx in np.linspace(-0.18, 0.14, 5):
            ax.plot([x+dx, x+dx+0.08], [-0.15, -0.22], 'k-', linewidth=0.5)


def _draw_moment_diagram_standard(ax, forces: Dict, model, scale: float, color: str):
    """
    绘制弯矩图（标准结构力学样式）
    - 梁：正弯矩向下（受拉侧），画抛物线形状
    - 柱：弯矩画在受拉侧
    - 标注端点和跨中数值
    """
    for elem_id, f in forces.items():
        if f.element_type == 'beam':
            start, end = model.beams[elem_id]
            x1, z1 = model.nodes[start]
            x2, z2 = model.nodes[end]
            
            # 梁弯矩图：端部为负弯矩(moment_min)，跨中为正弯矩(moment_max)
            # 均布荷载下呈抛物线分布
            L = (x2 - x1) / 1000  # 跨度(m)
            n_pts = 30  # 增加点数使曲线更光滑
            
            # 端部弯矩（负值，使用moment_min）和跨中弯矩（正值，使用moment_max）
            M_left = f.moment_min * scale   # 左端弯矩 (通常为负)
            M_right = f.moment_min * scale  # 右端弯矩 (假设对称)
            M_mid = f.moment_max * scale    # 跨中弯矩 (通常为正)
            
            # 使用三点抛物线插值: y = at² + bt + c
            # t=0: M_left, t=0.5: M_mid, t=1: M_right
            # 解出系数
            c = M_left
            a = 2 * (M_left + M_right - 2 * M_mid)
            b = M_right - M_left - a
            
            t = np.linspace(0, 1, n_pts)
            M_pts = a * t**2 + b * t + c
            
            x_pts = np.linspace(x1/1000, x2/1000, n_pts)
            z_base = z1/1000
            z_pts = z_base - M_pts  # 正弯矩向下绘制
            
            # 绘制弯矩图轮廓
            x_fill = np.concatenate([[x1/1000], x_pts, [x2/1000]])
            z_fill = np.concatenate([[z_base], z_pts, [z_base]])
            ax.fill(x_fill, z_fill, alpha=0.3, color=color, edgecolor=color, linewidth=1)
            ax.plot(x_pts, z_pts, color=color, linewidth=1.5)
            
            # 标注端点值
            # 注意：这里简化假设两端弯矩近似相等（均为负弯矩M_min）
            # 实际上外跨梁两端弯矩可能差异较大
            ax.annotate(f'{abs(f.moment_min):.0f}', 
                       xy=(x1/1000, z_base), xytext=(x1/1000-0.2, z_base+0.2),
                       fontsize=7, color=color)
            ax.annotate(f'{abs(f.moment_min):.0f}', 
                       xy=(x2/1000, z_base), xytext=(x2/1000+0.1, z_base+0.2),
                       fontsize=7, color=color)
            
            # 标注跨中弯矩
            x_mid = (x1 + x2) / 2000
            z_mid = z_base - M_mid
            ax.text(x_mid, z_mid - 0.3, f'{abs(f.moment_max):.0f}', 
                   ha='center', va='top', fontsize=7, color=color)
            # 标注跨中值
            ax.annotate(f'{f.M_design:.0f}', 
                       xy=((x1+x2)/2000, z_base - M_mid), 
                       xytext=((x1+x2)/2000, z_base - M_mid - 0.15),
                       fontsize=7, color=color, ha='center')
        
        else:  # 柱
            start, end = model.columns[elem_id]
            x1, z1 = model.nodes[start]
            x2, z2 = model.nodes[end]
            
            M = f.M_design * scale
            x_base = x1/1000
            
            # 柱弯矩图：假设线性分布
            ax.fill([x_base, x_base + M, x_base + M, x_base],
                   [z1/1000, z1/1000, z2/1000, z2/1000],
                   alpha=0.3, color=color, edgecolor=color, linewidth=1)
            ax.plot([x_base + M, x_base + M], [z1/1000, z2/1000], 
                   color=color, linewidth=1.5)
            
            # 标注底部值
            if f.M_design > 1:
                ax.annotate(f'{f.M_design:.0f}', 
                           xy=(x_base + M, z1/1000), 
                           xytext=(x_base + M + 0.1, z1/1000 + 0.1),
                           fontsize=6, color=color)


def _draw_shear_diagram_standard(ax, forces: Dict, model, scale: float, color: str):
    """
    绘制剪力图（标准结构力学样式）
    - 梁：正剪力向上，线性变化
    - 标注端点数值
    """
    for elem_id, f in forces.items():
        if f.element_type == 'beam':
            start, end = model.beams[elem_id]
            x1, z1 = model.nodes[start]
            x2, z2 = model.nodes[end]
            
            z_base = z1/1000
            V_left = f.shear_max * scale
            V_right = f.shear_min * scale
            
            # 剪力图：梯形（均布荷载下线性变化）
            ax.fill([x1/1000, x1/1000, x2/1000, x2/1000],
                   [z_base, z_base + V_left, z_base + V_right, z_base],
                   alpha=0.3, color=color, edgecolor=color, linewidth=1)
            ax.plot([x1/1000, x2/1000], [z_base + V_left, z_base + V_right], 
                   color=color, linewidth=1.5)
            
            # 标注端点值
            ax.annotate(f'{abs(f.shear_max):.0f}', 
                       xy=(x1/1000, z_base + V_left), 
                       xytext=(x1/1000-0.15, z_base + V_left),
                       fontsize=7, color=color, ha='right')
            ax.annotate(f'{abs(f.shear_min):.0f}', 
                       xy=(x2/1000, z_base + V_right), 
                       xytext=(x2/1000+0.15, z_base + V_right),
                       fontsize=7, color=color, ha='left')
        
        else:  # 柱
            start, end = model.columns[elem_id]
            x1, z1 = model.nodes[start]
            x2, z2 = model.nodes[end]
            
            V = f.V_design * scale
            x_base = x1/1000
            
            # 柱剪力：假设常量
            ax.fill([x_base, x_base + V, x_base + V, x_base],
                   [z1/1000, z1/1000, z2/1000, z2/1000],
                   alpha=0.3, color=color, edgecolor=color, linewidth=1)
            ax.plot([x_base + V, x_base + V], [z1/1000, z2/1000], 
                   color=color, linewidth=1.5)


def _draw_axial_diagram_standard(ax, forces: Dict, model, scale: float, color: str):
    """
    绘制轴力图（标准结构力学样式）
    - 柱：压力为正，画在构件侧面
    - 标注数值
    """
    for elem_id, f in forces.items():
        if f.element_type == 'column':
            start, end = model.columns[elem_id]
            x1, z1 = model.nodes[start]
            x2, z2 = model.nodes[end]
            
            N = f.N_design * scale
            x_base = x1/1000
            
            # 柱轴力：假设常量
            ax.fill([x_base, x_base + N, x_base + N, x_base],
                   [z1/1000, z1/1000, z2/1000, z2/1000],
                   alpha=0.3, color=color, edgecolor=color, linewidth=1)
            ax.plot([x_base + N, x_base + N], [z1/1000, z2/1000], 
                   color=color, linewidth=1.5)
            
            # 标注轴力值
            if f.N_design > 1:
                ax.annotate(f'{f.N_design:.0f}', 
                           xy=(x_base + N, (z1+z2)/2000), 
                           xytext=(x_base + N + 0.1, (z1+z2)/2000),
                           fontsize=6, color=color, ha='left', va='center')


# =============================================================================
# 收敛曲线
# =============================================================================

def plot_convergence(cost_history: List[float],
                     output_path: str = "收敛曲线.png") -> None:
    """
    绘制优化收敛曲线
    
    Args:
        cost_history: 每代最优造价历史
        output_path: 输出文件路径
    """
    fig, ax = plt.subplots(figsize=(10, 5))
    
    ax.plot(cost_history, 'b-', linewidth=2)
    ax.set_xlabel('迭代代数')
    ax.set_ylabel('最优造价 (元)')
    ax.set_title('遗传算法收敛曲线')
    ax.grid(True, alpha=0.3)
    
    # 标注初始和最终值
    ax.annotate(f'初始: {cost_history[0]:,.0f}元',
               xy=(0, cost_history[0]), xytext=(5, cost_history[0]),
               fontsize=10)
    ax.annotate(f'最终: {cost_history[-1]:,.0f}元',
               xy=(len(cost_history)-1, cost_history[-1]),
               fontsize=10)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    print(f"✓ 收敛曲线已保存: {output_path}")
    plt.close()


# =============================================================================
# 测试代码
# =============================================================================

if __name__ == "__main__":
    print("报表生成器模块测试")
    print("请运行 main.py 进行完整测试")
